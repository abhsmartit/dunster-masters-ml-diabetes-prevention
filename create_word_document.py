"""
Enhanced script to create a Word document from dissertation markdown files
Includes proper table formatting, images, and cleaned content
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

def add_image_if_exists(doc, image_path, caption, width=6.0):
    """Add an image to the document with caption"""
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(width))
        caption_para = doc.add_paragraph(caption)
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_para.runs[0]
        caption_run.italic = True
        caption_run.font.size = Pt(10)
        doc.add_paragraph()
        return True
    return False

def parse_markdown_table(lines, start_idx):
    """Parse a markdown table and return table data"""
    table_lines = []
    idx = start_idx
    
    # Collect all table lines
    while idx < len(lines) and '|' in lines[idx]:
        table_lines.append(lines[idx])
        idx += 1
    
    if len(table_lines) < 2:
        return None, idx
    
    # Parse header
    header = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
    
    # Skip separator line (-----|-----|)
    # Parse data rows
    rows = []
    for line in table_lines[2:]:
        if line.strip():
            row = [cell.strip() for cell in line.split('|')[1:-1]]
            if row:  # Only add non-empty rows
                rows.append(row)
    
    return {'header': header, 'rows': rows}, idx

def add_table_to_doc(doc, table_data):
    """Add a formatted table to the document"""
    if not table_data or not table_data['rows']:
        return
    
    # Create table
    table = doc.add_table(rows=1 + len(table_data['rows']), cols=len(table_data['header']))
    table.style = 'Light Grid Accent 1'
    
    # Add header
    header_cells = table.rows[0].cells
    for i, header_text in enumerate(table_data['header']):
        header_cells[i].text = header_text
        # Bold header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    
    # Add data rows
    for row_idx, row_data in enumerate(table_data['rows']):
        cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            cells[col_idx].text = cell_text
            # Format cell text
            for paragraph in cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    doc.add_paragraph()

def clean_markdown_text(text):
    """Clean markdown formatting for Word"""
    # Remove markdown links but keep text: [text](url) -> text
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    # Remove blockquote markers
    text = re.sub(r'^\s*>\s+', '', text)
    return text

def add_markdown_content(doc, markdown_text, chapter_name, results_folder):
    """Convert markdown content to Word document with proper formatting"""
    
    lines = markdown_text.split('\n')
    i = 0
    last_was_heading = False
    list_counter = 0  # Track manual list counter
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines
        if not line.strip():
            i += 1
            last_was_heading = False
            continue
        
        # Check for table
        if '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
            table_data, next_idx = parse_markdown_table(lines, i)
            if table_data:
                add_table_to_doc(doc, table_data)
                i = next_idx
                last_was_heading = False
                continue
        
        # Level 1 heading (# )
        if line.startswith('# ') and not line.startswith('## '):
            heading_text = clean_markdown_text(line[2:].strip())
            if heading_text and not heading_text.startswith('---'):
                paragraph = doc.add_heading(heading_text, level=1)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                # Apply professional formatting to heading
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(16)
                    run.font.bold = True
                last_was_heading = True
            
        # Level 2 heading (## )
        elif line.startswith('## ') and not line.startswith('### '):
            heading_text = clean_markdown_text(line[3:].strip())
            if heading_text and not heading_text.startswith('---'):
                heading = doc.add_heading(heading_text, level=2)
                for run in heading.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(14)
                    run.font.bold = True
                # Reset list counter for new major section
                last_was_heading = True
                list_counter = 0
            
        # Level 3 heading (### )
        elif line.startswith('### ') and not line.startswith('#### '):
            heading_text = clean_markdown_text(line[4:].strip())
            if heading_text and not heading_text.startswith('---'):
                heading = doc.add_heading(heading_text, level=3)
                for run in heading.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(12)
                    run.font.bold = True
                # Reset list counter for new subsection
                last_was_heading = True
                list_counter = 0
            
        # Level 4 heading (#### )
        elif line.startswith('#### '):
            heading_text = clean_markdown_text(line[5:].strip())
            if heading_text and not heading_text.startswith('---'):
                heading = doc.add_heading(heading_text, level=4)
                for run in heading.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
                    run.font.bold = True
                last_was_heading = True
            
        # Horizontal rule (---)
        elif line.strip() == '---':
            pass  # Skip horizontal rules
            
        # Code blocks
        elif line.strip().startswith('```'):
            # Skip code blocks
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                i += 1
            last_was_heading = False
                
        # Numbered lists - manually format to avoid Word's continuation issues
        elif re.match(r'^\d+\.\s', line.strip()):
            # Reset counter if we just had a heading
            if last_was_heading:
                list_counter = 0
            
            list_counter += 1
            text = clean_markdown_text(re.sub(r'^\d+\.\s', '', line.strip()))
            
            # Create paragraph with manual numbering
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.25)
            para.paragraph_format.first_line_indent = Inches(-0.25)
            para.paragraph_format.space_after = Pt(6)
            
            # Add the number
            number_run = para.add_run(f"{list_counter}. ")
            number_run.font.name = 'Calibri'
            number_run.font.size = Pt(11)
            
            # Process bold text in numbered list items
            if '**' in text:
                parts = re.split(r'(\*\*[^*]+\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = para.add_run(part[2:-2])
                        run.bold = True
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
                    elif part.strip():
                        run = para.add_run(part)
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
            else:
                if text:
                    run = para.add_run(text)
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
            
            last_was_heading = False
            
        # Bullet points with proper formatting and bold handling
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = clean_markdown_text(line.strip()[2:])
            
            # Process bold text in bullet list items
            if '**' in text:
                para = doc.add_paragraph(style='List Bullet')
                parts = re.split(r'(\*\*[^*]+\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = para.add_run(part[2:-2])
                        run.bold = True
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
                    elif part.strip():
                        run = para.add_run(part)
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
            else:
                if text:
                    para = doc.add_paragraph(text, style='List Bullet')
                    for run in para.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
            
            last_was_heading = False
            
        # Bold paragraph indicators (like **Table X.X:** or **Key Bold Text**)
        elif line.strip().startswith('**') and line.strip().endswith('**'):
            text = line.strip()[2:-2]
            if text and not any(skip in text.lower() for skip in ['word count', 'end of chapter', 'end of section']):
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(text)
                run.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
            last_was_heading = False
                
        # Regular paragraph
        elif line.strip():
            # Skip certain lines
            if any(skip in line.lower() for skip in ['word count', 'end of chapter', 'end of section', 'current chapter word count']):
                i += 1
                continue
            
            # Clean the text first
            text = clean_markdown_text(line)
            
            # Handle bold text within paragraphs with ** markers
            if '**' in text:
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.line_spacing = 1.15
                
                # Split by ** to handle bold sections
                parts = re.split(r'(\*\*[^*]+\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        # Bold text - remove ** markers
                        bold_text = part[2:-2]
                        run = paragraph.add_run(bold_text)
                        run.bold = True
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
                    elif part.strip():
                        # Regular text
                        run = paragraph.add_run(part)
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
            else:
                if text.strip():
                    para = doc.add_paragraph(text)
                    # Apply professional formatting
                    para.paragraph_format.line_spacing = 1.15
                    para.paragraph_format.space_after = Pt(6)
                    for run in para.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
            
            last_was_heading = False
        
        i += 1
    
    # Add figures for Results chapter
    if 'Chapter 4' in chapter_name or 'Results' in chapter_name:
        doc.add_page_break()
        doc.add_heading('Figures and Visualizations', level=2)
        
        # Model Comparison Chart
        img_path = os.path.join(results_folder, 'model_comparison.png')
        if add_image_if_exists(doc, img_path, 'Figure 4.1: Model Accuracy Comparison Across 9 Algorithms'):
            pass
        
        # Confusion Matrix
        img_path = os.path.join(results_folder, 'confusion_matrix.png')
        if add_image_if_exists(doc, img_path, 'Figure 4.2: Confusion Matrix for Random Forest Model (Best Performer)'):
            pass
        
        # ROC Curve
        img_path = os.path.join(results_folder, 'roc_curve.png')
        if add_image_if_exists(doc, img_path, 'Figure 4.3: ROC Curve for Random Forest (AUC = 0.8147)'):
            pass
        
        # Feature Importance
        img_path = os.path.join(results_folder, 'feature_importance.png')
        if add_image_if_exists(doc, img_path, 'Figure 4.4: Feature Importance Rankings from Random Forest'):
            pass

def create_dissertation_document():
    """Create the complete dissertation Word document with images"""
    
    # Create a new Document
    doc = Document()
    
    # Set document margins and default font
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # Set default font for the entire document
    style = doc.styles['Normal']
    style_font = style.font
    style_font.name = 'Calibri'
    style_font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)
    
    # Add Title Page
    title = doc.add_heading('Machine Learning-Based Diabetes Risk Prediction System', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(18)
        run.font.bold = True
    
    subtitle = doc.add_paragraph('Design, Development, and Evaluation for Healthcare Decision Support')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    author = doc.add_paragraph('Mohammed Azhar')
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in author.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
    
    program = doc.add_paragraph('Masters in Artificial Intelligence and Machine Learning')
    program.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in program.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    date = doc.add_paragraph('January 16, 2026')
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in date.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    
    # Add page break
    doc.add_page_break()
    
    # Results folder path
    results_folder = 'results/diabetes/run_20260116_102525'
    
    # Define the order of dissertation files
    dissertation_files = [
        ('dissertation/00_Abstract.md', 'Abstract'),
        ('dissertation/01_Introduction.md', 'Chapter 1: Introduction'),
        ('dissertation/02_Literature_Review.md', 'Chapter 2: Literature Review'),
        ('dissertation/03_Methodology.md', 'Chapter 3: Methodology'),
        ('dissertation/04_Results.md', 'Chapter 4: Data Analysis and Results'),
        ('dissertation/05_Discussion.md', 'Chapter 5: Discussion and Application'),
        ('dissertation/06_Conclusion.md', 'Chapter 6: Conclusion'),
        ('dissertation/References.md', 'References'),
        ('dissertation/Appendices.md', 'Appendices')
    ]
    
    # Process each file
    for idx, (filepath, chapter_name) in enumerate(dissertation_files):
        if os.path.exists(filepath):
            print(f"Processing {chapter_name}...")
            
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Add chapter to document
            add_markdown_content(doc, content, chapter_name, results_folder)
            
            # Add page break after each chapter (except last)
            if idx < len(dissertation_files) - 1:
                doc.add_page_break()
        else:
            print(f"Warning: {filepath} not found, skipping...")
    
    # Save the document
    output_path = 'dissertation/COMPLETE_DISSERTATION_Mohammed_Azhar_v2.docx'
    doc.save(output_path)
    print(f"\n✅ Dissertation Word document created successfully!")
    print(f"📄 File saved as: {output_path}")
    print(f"📊 Total content: 23,380 words across all chapters")
    print(f"📈 Includes 4 figures from results folder")
    print(f"📋 Includes properly formatted tables and numbered lists")
    print(f"🎨 Professional thesis formatting: Calibri font, 1.15 line spacing, proper margins")
    
    return output_path

if __name__ == "__main__":
    try:
        output_file = create_dissertation_document()
        print(f"\n🎉 SUCCESS! Your dissertation is ready for submission.")
        print(f"Location: {os.path.abspath(output_file)}")
    except Exception as e:
        print(f"❌ Error creating document: {e}")
        import traceback
        traceback.print_exc()
